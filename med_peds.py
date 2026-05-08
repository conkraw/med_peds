# app.py
import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import requests

TEMPLATE_URLS = {
    "Mid-Year - No Concerns": "https://raw.githubusercontent.com/conkraw/med_peds/main/template/my_pgy1_to_4_no_concerns.docx",
    "End-of-Year - No Concerns": "https://raw.githubusercontent.com/conkraw/med_peds/main/template/eoy_pgy1_to_3_no_concerns.docx",
    "Mid-Year - Concerns": "https://raw.githubusercontent.com/conkraw/med_peds/main/template/my_concerns.docx",
    "End-of-Year_PGY4 - No Concerns": "https://raw.githubusercontent.com/conkraw/med_peds/main/template/eoy_pgy4.docx",
}

st.set_page_config(page_title="CCC Letter Generator", layout="centered")

TODAY_STR = date.today().strftime("%B %d, %Y").replace(" 0", " ")

st.title("CCC Letter Generator")
st.caption(f"Today's date: {TODAY_STR}")

# ----------------------------
# Helper functions
# ----------------------------
def replace_simple_text(
    paragraph,
    replacements: dict,
    font_name: str = "Times New Roman",
    font_size_pt: int = 11,
):
    if not paragraph.text:
        return

    full_text = paragraph.text
    new_text = full_text

    for old, new in replacements.items():
        new_text = new_text.replace(old, new)

    if new_text == full_text:
        return

    paragraph.clear()
    run = paragraph.add_run(new_text)

    # Force font + size
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)

    # Ensure Word actually honors the font
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def replace_everywhere(doc: Document, replacements: dict):
    for p in doc.paragraphs:
        replace_simple_text(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_simple_text(p, replacements)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_simple_text(p, replacements)
        for p in section.footer.paragraphs:
            replace_simple_text(p, replacements)

def insert_date_at_top(doc: Document, date_line: str):
    if doc.paragraphs:
        doc.paragraphs[0].insert_paragraph_before(date_line)
    else:
        doc.add_paragraph(date_line)

def replace_ccc_text_block(
    doc: Document,
    replacement_text: str,
    token: str = "CCC_TEXT",
    font_name: str = "Times New Roman",
    font_size_pt: int = 11,
):
    """
    Replace the paragraph containing `token` with multiple paragraphs,
    formatted as Times New Roman, 11 pt.
    Preserves all surrounding content.
    """

    blocks = [line for line in replacement_text.split("\n") if line.strip()]

    def _insert_formatted_paragraph(before_paragraph, text):
        new_p = before_paragraph.insert_paragraph_before()
        run = new_p.add_run(text)

        # Font settings
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)

        # Ensure Word respects the font (important!)
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        run._element.rPr.rFonts.set(qn("w:cs"), font_name)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _process_paragraph(p):
        if token not in p.text:
            return False

        for block in blocks:
            _insert_formatted_paragraph(p, block)

        # Remove ONLY the placeholder paragraph
        parent = p._p.getparent()
        parent.remove(p._p)
        return True

    # Body paragraphs
    for p in list(doc.paragraphs):
        if _process_paragraph(p):
            return True

    # Tables (robustness)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    if _process_paragraph(p):
                        return True

    return False

# ----------------------------
# UI
# ----------------------------
with st.form("inputs"):
    name = st.text_input("Person's name (e.g., Jane Smith)")
    month = st.selectbox(
        "Month",
        ["January","February","March","April","May","June","July","August",
         "September","October","November","December"],
        index=date.today().month - 1,
    )
    year = st.number_input("Year", min_value=2000, max_value=2100,
                           value=date.today().year, step=1)

    with st.expander("Show AI Prompt Helper"):
        st.code(
            '''Write a concise Clinical Competency Committee (CCC) summary paragraph in a supportive, professional tone appropriate for both resident feedback documentation and inclusion in a warm letter of support.

Formatting and style requirements:

Write as a single cohesive 3-sentence paragraph.
Do not use the word “resident.”
Begin with: “Across both Internal Medicine and Pediatrics settings,”
Use polished narrative language similar to faculty evaluation summaries.
Emphasize longitudinal growth, strengths, and responsiveness to feedback.
Present strengths first, followed by developmental growth areas.
Frame growth areas constructively and developmentally rather than critically.
Avoid bullet points, headings, or overly harsh language.
Highlight trajectory of improvement over time.
End with a future-oriented statement such as “expected to continue strengthening with experience and coaching” or similar language.
Maintain a balanced, encouraging, and professional tone throughout.

Content to incorporate:
[INSERT CONTENT HERE]''',
            language="text"
        )
    ccc_text = st.text_area(
        "CCC personalized paragraph(s)",
        height=220,
        help="Line breaks will be preserved as separate paragraphs in Word."
    )

    template_choice = st.selectbox("Select template",["Mid-Year - No Concerns", "End-of-Year - No Concerns", "Mid-Year - Concerns", "End-of-Year_PGY4 - No Concerns"])
    
    submitted = st.form_submit_button("Generate Word Document")

# ----------------------------
# Generate document
# ----------------------------
if submitted:
    if not name.strip() or not ccc_text.strip():
        st.error("Please complete all required fields.")
        st.stop()

    month_year = f"{month} {int(year)}"

    template_url = TEMPLATE_URLS[template_choice]
    response = requests.get(template_url)
    response.raise_for_status()
    doc = Document(BytesIO(response.content))

    # Simple placeholder replacements
    replace_everywhere(
        doc,
        {
            "xxx": name.strip(),
            "Date": TODAY_STR,
            "Month of Year": month_year,
            "Month of  Year": month_year,
        }
    )

    # Replace CCC_TEXT block with multi-paragraph input
    replace_ccc_text_block(doc, ccc_text)

    # Output
    out = BytesIO()
    doc.save(out)
    out.seek(0)

    safe_name = name.strip().replace(" ", "_")
    filename = f"CCC_Letter_{safe_name}.docx"

    st.success("Document created successfully.")
    st.download_button(
        "Download Word Document",
        data=out.getvalue(),
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

